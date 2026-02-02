#!/usr/bin/env python3
from pathlib import Path
from docx import Document

TEMPLATES_DIR = Path("templates")      # folder with your templates
FIXED_DIR = TEMPLATES_DIR / "fixed"    # folder for updated templates
FIXED_DIR.mkdir(exist_ok=True)

# --- Mapping of old variable names -> new variable names ---
REPLACEMENTS = {
    "defendantpluralUPPER": "defendantplural",
    "venuecounty": "venue",
    "Plaintiffname": "plaintiffname",
    "countyjurUPPER": "jurisdiction",
    "venuecountyUPPER": "venue",
    "defendantpluralUPPER": "defendantplural",
    "affirmdefplrlCAP": "affirmdefplrl",
    "Defendantname": "defendantname",
    # add all other substitutions here
}

def replace_in_paragraph(paragraph, replacements):
    """Replaces all old variables in a paragraph, even if Word split them into multiple runs."""
    full_text = "".join(run.text for run in paragraph.runs)
    for old_var, new_var in replacements.items():
        full_text = full_text.replace(f"{{{{{old_var}}}}}", f"{{{{{new_var}}}}}")
    # Clear runs and set new text
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)

def replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, replacements)

# --- Process all .docx templates ---
for file_path in TEMPLATES_DIR.glob("*.docx"):
    print(f"Processing {file_path.name}...")
    doc = Document(file_path)
    for p in doc.paragraphs:
        replace_in_paragraph(p, REPLACEMENTS)
    for table in doc.tables:
        replace_in_table(table, REPLACEMENTS)
    doc.save(FIXED_DIR / file_path.name)

print(f"All templates updated. Fixed versions saved in '{FIXED_DIR}'.")
