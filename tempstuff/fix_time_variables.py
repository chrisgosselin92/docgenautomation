#!/usr/bin/env python3
from pathlib import Path
from docx import Document

TEMPLATES_DIR = Path("templates")          # source templates
FIXED_DIR = TEMPLATES_DIR / "fixed_time"   # output folder
FIXED_DIR.mkdir(exist_ok=True)

# --- Time variables to convert {{var}} -> [[var]] ---
TIME_VARIABLES = {
    "currentdate",
    "current",
    "currentday",
    "currentmonth",
    "year",
}

def fix_time_vars_in_paragraph(paragraph):
    """
    Replace {{timevar}} with [[timevar]] in a paragraph,
    safely handling Word run splitting.
    """
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)

    for var in TIME_VARIABLES:
        full_text = full_text.replace(
            f"{{{{{var}}}}}",
            f"[[{var}]]"
        )

    # Rewrite paragraph cleanly
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = full_text

def fix_time_vars_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                fix_time_vars_in_paragraph(paragraph)

# --- Process all templates ---
for file_path in TEMPLATES_DIR.glob("*.docx"):
    print(f"Processing {file_path.name}...")
    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        fix_time_vars_in_paragraph(paragraph)

    for table in doc.tables:
        fix_time_vars_in_table(table)

    doc.save(FIXED_DIR / file_path.name)

print(f"Done. Time variables converted and saved to '{FIXED_DIR}'.")
