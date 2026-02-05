# modules/bracket_variables.py
"""
Handles [[variable]] replacement for dynamic Excel content.
These variables are embedded in Excel responses and replaced after list creation.
"""

import re
from docx import Document
from modules.db import get_variables


def extract_bracket_variables(template_path):
    """
    Extract all [[variable]] patterns from document.
    Returns set of variable names.
    """
    doc = Document(template_path)
    bracket_vars = set()
    
    pattern = r'\[\[([a-zA-Z_][a-zA-Z0-9_]*)\]\]'
    
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        for var_name in matches:
            bracket_vars.add(var_name)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    for var_name in matches:
                        bracket_vars.add(var_name)
    
    return bracket_vars


def replace_bracket_variables(doc_path, client_id):
    """
    Replace [[variable]] with values from client database.
    Also handles grammar variables like [[he_she_they]].
    """
    from modules.db import get_variables
    from modules.grammar import GRAMMAR_RULES
    
    doc = Document(doc_path)
    
    # Get all client variables
    client_vars = get_variables("client", client_id)
    
    # Get grammar settings from client
    count = "singular"  # Default
    gender = "male"     # Default
    
    # Check if client has defendant_count and gender
    defendant_count = client_vars.get("defendant_count", "1")
    try:
        if int(defendant_count) > 1:
            count = "plural"
    except:
        pass
    
    client_gender = client_vars.get("gender", "male").lower()
    if client_gender in ["male", "female"]:
        gender = client_gender
    
    def get_replacement(var_name):
        """Get the replacement value for a [[variable]]"""
        # Check if it's a grammar variable first
        if var_name in GRAMMAR_RULES:
            rule = GRAMMAR_RULES[var_name]
            
            # Simple singular/plural rules
            if "singular" in rule and "plural" in rule:
                return rule[count]
            
            # Gender-dependent pronoun rules
            if count == "plural":
                return rule["plural"]
            else:
                key = f"singular_{gender}"
                return rule.get(key, rule.get("singular", f"[[{var_name}]]"))
        
        # Not a grammar variable - get from client database
        value = client_vars.get(var_name, "")
        if not value:
            # Check for common derived variables
            if var_name == "plaintiff":
                return client_vars.get("clientname", client_vars.get("firstname", "")) + " " + client_vars.get("lastname", "")
            elif var_name == "defendant":
                return client_vars.get("defendantname", "Defendant")
        
        return value if value else f"[[{var_name}]]"
    
    def replace_in_paragraph(paragraph):
        """Replace [[variable]] in a paragraph"""
        pattern = r'\[\[([a-zA-Z_][a-zA-Z0-9_]*)\]\]'
        matches = re.findall(pattern, paragraph.text)
        
        for var_name in matches:
            replacement = get_replacement(var_name)
            original_pattern = f"[[{var_name}]]"
            
            for run in paragraph.runs:
                if original_pattern in run.text:
                    run.text = run.text.replace(original_pattern, str(replacement))
    
    # Replace in all document parts
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph)
    
    doc.save(doc_path)