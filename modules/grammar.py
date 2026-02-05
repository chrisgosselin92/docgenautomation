# modules/grammar.py
"""
Grammar and plurality handling for document generation.
Handles (@variable@) patterns for automatic grammatical agreement.
"""

import tkinter as tk
from tkinter import messagebox
import re
from docx import Document

# Grammar rules based on client count and gender
GRAMMAR_RULES = {
    # Basic plurality
    "clients": {"singular": "client", "plural": "clients"},
    "plural_s": {"singular": "", "plural": "s"},  # jump/jumps
    "standard-s_plural": {"singular": "s", "plural": ""},  # states/state
    "plural_ies": {"singular": "ies", "plural": "y"},  # denies/deny
    
    # Legal party terms
    "defendant_defendants": {"singular": "defendant", "plural": "defendants"},
    "plaintiff_plaintiffs": {"singular": "plaintiff", "plural": "plaintiffs"},
    "party_parties": {"singular": "party", "plural": "parties"},
    "petitioner_petitioners": {"singular": "petitioner", "plural": "petitioners"},
    "respondent_respondents": {"singular": "respondent", "plural": "respondents"},
    "movant_movants": {"singular": "movant", "plural": "movants"},
    "appellee_appellees": {"singular": "appellee", "plural": "appellees"},
    "appellant_appellants": {"singular": "appellant", "plural": "appellants"},
    
    # Verb conjugation (basic)
    "is_are": {"singular": "is", "plural": "are"},
    "was_were": {"singular": "was", "plural": "were"},
    "have_has": {"singular": "has", "plural": "have"},
    "does_do": {"singular": "does", "plural": "do"},
    
    # Verb conjugation (legal/specific)
    "believes_believe": {"singular": "believes", "plural": "believe"},
    "denies_deny": {"singular": "denies", "plural": "deny"},
    "alleges_allege": {"singular": "alleges", "plural": "allege"},
    "requests_request": {"singular": "requests", "plural": "request"},
    "moves_move": {"singular": "moves", "plural": "move"},
    "opposes_oppose": {"singular": "opposes", "plural": "oppose"},
    
    # Subject pronouns (gender-based for singular, "they" for plural)
    "he_she_they": {
        "singular_male": "he",
        "singular_female": "she",
        "plural": "they"
    },
    
    # Object pronouns
    "him_her_them": {
        "singular_male": "him",
        "singular_female": "her",
        "plural": "them"
    },
    
    # Possessive adjectives
    "his_her_their": {
        "singular_male": "his",
        "singular_female": "her",
        "plural": "their"
    },
    
    # Possessive pronouns
    "his_hers_theirs": {
        "singular_male": "his",
        "singular_female": "hers",
        "plural": "theirs"
    },
    
    # Reflexive pronouns
    "himself_herself_themselves": {
        "singular_male": "himself",
        "singular_female": "herself",
        "plural": "themselves"
    },
}


def prompt_grammar_settings(parent):
    """
    Prompt user for grammar settings: singular/plural and gender.
    Returns dict with 'count' and 'gender' keys.
    """
    dialog = tk.Toplevel(parent)
    dialog.title("Grammar Settings")
    dialog.geometry("550x450")
    dialog.grab_set()
    
    tk.Label(dialog, text="Grammar Settings for Document", font=("Arial", 14, "bold")).pack(pady=20)
    
    # Client count (singular/plural)
    tk.Label(dialog, text="Number of clients/parties in this matter:", font=("Arial", 11)).pack(pady=10)
    count_var = tk.StringVar(value="singular")
    
    count_frame = tk.Frame(dialog)
    count_frame.pack(pady=5)
    tk.Radiobutton(count_frame, text="Single client/party", variable=count_var, value="singular", font=("Arial", 10)).pack(anchor="w", padx=20)
    tk.Label(count_frame, text="(uses: he/she, him/her, his/hers, defendant, party)", fg="gray", font=("Arial", 8)).pack(anchor="w", padx=40)
    
    tk.Radiobutton(count_frame, text="Multiple clients/parties", variable=count_var, value="plural", font=("Arial", 10)).pack(anchor="w", padx=20, pady=(10,0))
    tk.Label(count_frame, text="(uses: they, them, their, defendants, parties)", fg="gray", font=("Arial", 8)).pack(anchor="w", padx=40)
    
    # Gender (only relevant for singular)
    tk.Label(dialog, text="Gender (for singular client only):", font=("Arial", 11)).pack(pady=(20, 10))
    gender_var = tk.StringVar(value="male")
    
    gender_frame = tk.Frame(dialog)
    gender_frame.pack(pady=5)
    tk.Radiobutton(gender_frame, text="Male (he/him/his/himself)", variable=gender_var, value="male", font=("Arial", 10)).pack(anchor="w", padx=20)
    tk.Radiobutton(gender_frame, text="Female (she/her/hers/herself)", variable=gender_var, value="female", font=("Arial", 10)).pack(anchor="w", padx=20)
    
    result = {"count": None, "gender": None}
    
    def on_submit():
        result["count"] = count_var.get()
        result["gender"] = gender_var.get()
        dialog.destroy()
    
    tk.Button(dialog, text="Continue", command=on_submit, width=15, bg="#4CAF50", fg="white", font=("Arial", 11)).pack(pady=20)
    
    dialog.wait_window()
    
    return result


def extract_grammar_variables(template_path):
    """
    Extract all (@variable@) grammar patterns from template.
    Returns set of variable names.
    """
    doc = Document(template_path)
    grammar_vars = set()
    
    pattern = r'\(@([a-zA-Z_][a-zA-Z0-9_-]*?)@\)'
    
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        for var_name in matches:
            grammar_vars.add(var_name)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    for var_name in matches:
                        grammar_vars.add(var_name)
    
    return grammar_vars


def replace_grammar_variables(doc_path, grammar_settings):
    """
    Replace (@variable@) with appropriate grammatical forms.
    """
    doc = Document(doc_path)
    count = grammar_settings["count"]  # "singular" or "plural"
    gender = grammar_settings["gender"]  # "male" or "female"
    
    def get_replacement(var_name):
        """Get the replacement text for a grammar variable"""
        if var_name not in GRAMMAR_RULES:
            return f"(@{var_name}@)"  # Leave unchanged if unknown
        
        rule = GRAMMAR_RULES[var_name]
        
        # Simple singular/plural rules
        if "singular" in rule and "plural" in rule:
            return rule[count]
        
        # Gender-dependent pronoun rules
        if count == "plural":
            return rule["plural"]
        else:
            # Singular - check gender
            key = f"singular_{gender}"
            return rule.get(key, rule.get("singular", f"(@{var_name}@)"))
    
    def replace_in_paragraph(paragraph):
        pattern = r'\(@([a-zA-Z_][a-zA-Z0-9_-]*?)@\)'
        matches = re.findall(pattern, paragraph.text)
        
        for var_name in matches:
            replacement = get_replacement(var_name)
            original_pattern = f"(@{var_name}@)"
            
            for run in paragraph.runs:
                if original_pattern in run.text:
                    run.text = run.text.replace(original_pattern, replacement)
    
    # Replace in all document parts
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph)
    
    doc.save(doc_path)