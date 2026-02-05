# modules/docgen.py
"""
Comprehensive document generation module with support for 5 variable types:
1. Standard variables - prompt for missing values, store in DB
2. Date variables - auto-fill with ordinal formatting
3. Concatenated/Derived variables - composite or conditional variables
4. Dynamic variables - from Excel with <<>> delimiters
5. Opposing Counsel variables - from attorney database with (()) delimiters
"""

import tkinter as tk
from tkinter import messagebox, ttk, simpledialog
from pathlib import Path
from datetime import datetime
import re
import pandas as pd
from tkinter import simpledialog
from modules.bracket_variables import extract_bracket_variables, replace_bracket_variables
from modules.grammar import extract_grammar_variables, replace_grammar_variables, prompt_grammar_settings
from docxtpl import DocxTemplate
from docx import Document
from modules.db import (
    DB_PATH,
    list_clients,
    get_variables,
    set_variable,
    get_variable_value_for_client,
    variable_exists,
    set_variable_meta,
    list_all_concats,
    get_all_variables_for_client
)

# Try to import improved version first, fall back to original
try:
    from modules.editconcatvariable_improved import open_concat_editor, get_or_build_derived_value
except ImportError:
    from modules.editconcatvariable import open_concat_editor, get_or_build_derived_value


# =============================================================================
# VARIABLE TYPE 1: DATE VARIABLES - Auto-filled system dates
# =============================================================================

def get_system_date_context():
    """
    Returns a dictionary of date-related variables with proper formatting.
    Includes ordinal formatting (1st, 2nd, 3rd, etc.)
    """
    now = datetime.now()
    
    def ordinal(n):
        """Convert number to ordinal string (1st, 2nd, 3rd, etc.)"""
        if 10 <= n % 100 <= 20:
            suffix = "th"
        else:
            suffix = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
        return f"{n}{suffix}"
    
    return {
        # Day variations
        "currentday": ordinal(now.day),  # "2nd"
        "currentdaynum": str(now.day),    # "2"
        "currentdayordinal": ordinal(now.day),  # "2nd"
        
        # Month variations
        "currentmonth": now.strftime("%B"),  # "February"
        "monthabbr": now.strftime("%b"),     # "Feb"
        "monthnum": str(now.month),          # "2"
        
        # Year variations
        "year": str(now.year),        # "2026"
        "year2": now.strftime("%y"),  # "26"
        
        # Weekday variations
        "weekday": now.strftime("%A"),      # "Monday"
        "weekdayabbr": now.strftime("%a"),  # "Mon"
        
        # Common combinations
        "today": now.strftime("%B %d, %Y"),  # "February 02, 2026"
        "todayshort": now.strftime("%m/%d/%Y"),  # "02/02/2026"
    }


# =============================================================================
# VARIABLE TYPE 2: STANDARD VARIABLES - Prompt and store in DB
# =============================================================================

def prompt_for_variable(parent, var_name, client_id, all_client_vars, allow_cancel=True):
    """
    Prompts user to enter a value for a missing variable.
    Stores the value in the database for this client.
    """
    # Create a simple input dialog
    dialog = tk.Toplevel(parent)
    dialog.title(f"Enter Value for {var_name}")
    dialog.geometry("500x200")
    dialog.grab_set()
    
    tk.Label(dialog, text=f"Variable: {var_name}", font=("Arial", 12, "bold")).pack(pady=10)
    tk.Label(dialog, text="This variable is missing for this client.").pack(pady=5)
    
    value_var = tk.StringVar()
    tk.Label(dialog, text="Enter value:").pack(pady=5)
    entry = tk.Entry(dialog, textvariable=value_var, width=50)
    entry.pack(pady=5)
    entry.focus()
    
    result = {"value": None, "cancelled": False}
    
    def on_submit():
        val = value_var.get().strip()
        if not val:
            messagebox.showwarning("Empty Value", "Please enter a value or cancel.", parent=dialog)
            return
        result["value"] = val
        dialog.destroy()
    
    def on_skip():
        result["value"] = ""
        dialog.destroy()
    
    def on_cancel():
        result["cancelled"] = True
        dialog.destroy()
    
    button_frame = tk.Frame(dialog)
    button_frame.pack(pady=10)
    tk.Button(button_frame, text="Submit", command=on_submit, width=10).pack(side="left", padx=5)
    tk.Button(button_frame, text="Skip (Leave Empty)", command=on_skip, width=15).pack(side="left", padx=5)
    
    if allow_cancel:
        tk.Button(button_frame, text="Cancel Generation", command=on_cancel, width=15, bg="#f44336", fg="white").pack(side="left", padx=5)
    
    # Allow Enter key to submit
    entry.bind("<Return>", lambda e: on_submit())
    
    dialog.wait_window()
    
    if result["cancelled"]:
        raise Exception("User cancelled document generation")
    
    if result["value"] is not None:
        # Store in database
        set_variable("client", client_id, var_name, result["value"])
        # Also add to metadata if it doesn't exist
        if not variable_exists(var_name):
            set_variable_meta(var_name, var_type="string", description=f"User-defined: {var_name}")
    
    return result["value"] or ""


# =============================================================================
# VARIABLE TYPE 3: CONCATENATED & DERIVED VARIABLES
# =============================================================================

def handle_concatenated_variable(parent, var_name, client_id, all_client_vars):
    """
    Handles concatenated variables (combinations of other variables).
    Opens GUI with pre-filled variable name if not already defined.
    """
    # Check if this is already defined as a concat variable
    concats = {c["var_name"]: c for c in list_all_concats()}
    
    if var_name in concats:
        # Build the value from components
        comp_list = concats[var_name]["components"]
        sep = concats[var_name].get("separator", " ")
        values = []
        for comp in comp_list:
            val = all_client_vars.get(comp, "")
            values.append(str(val))
        return sep.join(values)
    
    # Not defined - ask user to build it
    response = messagebox.askyesno(
        "Concatenated Variable",
        f"Variable '{var_name}' appears to be a concatenated variable.\n\n"
        f"Would you like to build it from existing variables?",
        parent=parent
    )
    
    if response:
        # Open concat editor WITH PRE-FILLED NAME
        result = open_concat_editor(parent, entity_type="client", prefill_name=var_name)
        
        # If result is returned directly, use it
        if result:
            return result
        
        # Try again after editor closes
        concats = {c["var_name"]: c for c in list_all_concats()}
        if var_name in concats:
            comp_list = concats[var_name]["components"]
            sep = concats[var_name].get("separator", " ")
            values = []
            for comp in comp_list:
                val = all_client_vars.get(comp, "")
                values.append(str(val))
            return sep.join(values)
    
    # Fallback to manual entry
    return prompt_for_variable(parent, var_name, client_id, all_client_vars)


def handle_derived_variable(parent, var_name, client_id, all_client_vars):
    """
    Handles derived variables (conditional/grammatical transformations).
    Examples: pluralization, pronouns based on gender, verb conjugation, etc.
    """
    base_var = var_name
    transformation = None
    
    # Detect transformation patterns
    if var_name.endswith("_plural"):
        base_var = var_name.replace("_plural", "")
        transformation = "plural"
    elif var_name.endswith("_possessive"):
        base_var = var_name.replace("_possessive", "")
        transformation = "possessive"
    elif var_name in ["he_she", "him_her", "his_her", "his_hers", "He_She", "Him_Her", "His_Her", "His_Hers"]:
        transformation = "pronoun"
        base_var = "gender"
    elif "_deny" in var_name or "_denies" in var_name:
        transformation = "verb_conjugate_deny"
        base_var = var_name.split("_")[0]  # Get base (e.g., "defendant" from "defendant_deny")
    
    # PLURALIZATION
    if transformation == "plural":
        base_value = all_client_vars.get(base_var, "")
        if base_value:
            # English pluralization rules
            lower = base_value.lower()
            if lower.endswith(("s", "ss", "x", "z", "ch", "sh")):
                return base_value + "es"
            elif lower.endswith("y") and len(base_value) > 1 and base_value[-2] not in "aeiou":
                return base_value[:-1] + "ies"
            elif lower.endswith("f"):
                return base_value[:-1] + "ves"
            elif lower.endswith("fe"):
                return base_value[:-2] + "ves"
            elif lower.endswith("o") and len(base_value) > 1 and base_value[-2] not in "aeiou":
                return base_value + "es"
            else:
                return base_value + "s"
    
    # POSSESSIVE
    elif transformation == "possessive":
        base_value = all_client_vars.get(base_var, "")
        if base_value:
            if base_value.endswith("s"):
                return base_value + "'"
            else:
                return base_value + "'s"
    
    # PRONOUNS
    elif transformation == "pronoun":
        gender = all_client_vars.get("gender", "").lower()
        is_capitalized = var_name[0].isupper()
        
        pronoun_map = {
            "he_she": {"male": "he", "female": "she", "m": "he", "f": "she", "other": "they"},
            "him_her": {"male": "him", "female": "her", "m": "him", "f": "her", "other": "them"},
            "his_her": {"male": "his", "female": "her", "m": "his", "f": "her", "other": "their"},
            "his_hers": {"male": "his", "female": "hers", "m": "his", "f": "hers", "other": "theirs"},
        }
        
        key = var_name.lower()
        if key in pronoun_map and gender in pronoun_map[key]:
            result = pronoun_map[key][gender]
            return result.capitalize() if is_capitalized else result
    
    # VERB CONJUGATION (deny/denies based on count)
    elif transformation == "verb_conjugate_deny":
        count = int(all_client_vars.get("defendant_count", 1))
        if count == 1:
            return "denies"
        else:
            return "deny"
    
    # If no transformation matched, treat as concatenated
    return handle_concatenated_variable(parent, var_name, client_id, all_client_vars)


# =============================================================================
# VARIABLE TYPE 4: DYNAMIC VARIABLES (from Excel with <<>> delimiters)
# =============================================================================

def extract_dynamic_variables_from_template(template_path):
    """
    Extracts dynamic variables marked with <<variable_name>> from a template.
    Returns a set of variable names.
    """
    doc = Document(template_path)
    dynamic_vars = set()
    
    # Pattern to match <<variable>> or <<variable_modifier>>
    pattern = r'<<([a-zA-Z_][a-zA-Z0-9_]*)(?:_([a-zA-Z]+))?>>'
    
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        for var_name, modifier in matches:
            dynamic_vars.add((var_name, modifier if modifier else None))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    for var_name, modifier in matches:
                        dynamic_vars.add((var_name, modifier if modifier else None))
    
    return dynamic_vars


def prompt_dynamic_variable_from_excel(parent, var_name, client_id, excel_path="dynamicpleadingresponses.xlsx"):
    """
    Prompts user to select value for a dynamic variable from Excel sheet.
    Column A: Selection options (what user sees)
    Column B: Output values (what gets inserted)
    Column C: Instructions/question for user
    Column D1: "TRUE" if single-use, "FALSE" for multi-entry numbered list
    """
    try:
        df = pd.read_excel(excel_path, sheet_name=var_name)
    except Exception as e:
        return None
    
    if df.empty or len(df.columns) < 2:
        return None
    
    # Check column D1 for TRUE/FALSE (index 3)
    is_single_use = True
    use_numbered_list = False
    
    if len(df.columns) >= 4:
        try:
            d1_value = df.iloc[0, 3]
            
            if pd.isna(d1_value):
                d1_str = "TRUE"
            elif isinstance(d1_value, bool):
                d1_str = "TRUE" if d1_value else "FALSE"
            else:
                d1_str = str(d1_value).strip().upper()
            
            if d1_str == "FALSE":
                is_single_use = False
                use_numbered_list = True
            else:
                is_single_use = True
                use_numbered_list = False
        except Exception as e:
            print(f"Error reading D1: {e}, defaulting to TRUE (single-use)")
            is_single_use = True
            use_numbered_list = False
    
    # Get options from columns A and B
    options = []
    for idx, row in df.iterrows():
        try:
            display = str(row.iloc[0]).strip()
            output = str(row.iloc[1]).strip()
            if display and output and display != 'nan' and output != 'nan':
                options.append((display, output))
        except:
            continue
    
    if not options:
        return None
    
    # ========================================================================
    # FALSE = MULTI-ENTRY NUMBERED LIST (Loop until "This is last paragraph")
    # ========================================================================
    if use_numbered_list:
        collected_items = []
        entry_num = 1
        
        while True:
            # Create a NEW dialog for EACH entry
            dialog = tk.Toplevel(parent)
            dialog.title(f"{var_name} - Entry #{entry_num}")
            dialog.geometry("750x600")
            dialog.grab_set()
            
            # Header
            header_frame = tk.Frame(dialog, bg="#2196F3", height=60)
            header_frame.pack(fill="x")
            header_frame.pack_propagate(False)
            
            tk.Label(
                header_frame,
                text=f"Building Numbered List: {var_name}",
                font=("Arial", 14, "bold"),
                bg="#2196F3",
                fg="white"
            ).pack(pady=10)
            
            tk.Label(
                header_frame,
                text=f"Entry #{entry_num}",
                font=("Arial", 11),
                bg="#2196F3",
                fg="white"
            ).pack()
            
            # Show previously collected items
            if collected_items:
                collected_frame = tk.Frame(dialog, bg="#f0f0f0", relief="sunken", bd=2)
                collected_frame.pack(fill="x", padx=10, pady=10)
                
                tk.Label(
                    collected_frame,
                    text=f"✓ Items added so far ({len(collected_items)}):",
                    font=("Arial", 10, "bold"),
                    bg="#f0f0f0"
                ).pack(anchor="w", padx=5, pady=5)
                
                # Show last 3 items to save space
                display_items = collected_items[-3:] if len(collected_items) > 3 else collected_items
                start_num = len(collected_items) - len(display_items) + 1
                
                if len(collected_items) > 3:
                    tk.Label(
                        collected_frame,
                        text=f"... ({len(collected_items) - 3} earlier items)",
                        bg="#f0f0f0",
                        fg="gray",
                        font=("Arial", 8, "italic")
                    ).pack(anchor="w", padx=10)
                
                for i, item in enumerate(display_items, start=start_num):
                    item_text = item if len(item) <= 70 else item[:70] + "..."
                    tk.Label(
                        collected_frame,
                        text=f"{i}. {item_text}",
                        bg="#f0f0f0",
                        anchor="w",
                        font=("Arial", 9)
                    ).pack(anchor="w", padx=10, pady=1)
            
            # Selection area
            select_frame = tk.Frame(dialog)
            select_frame.pack(fill="both", expand=True, padx=10, pady=10)
            
            tk.Label(
                select_frame,
                text=f"Select item #{entry_num}:",
                font=("Arial", 11, "bold")
            ).pack(anchor="w", pady=(0, 5))
            
            # Search box
            search_var = tk.StringVar()
            search_container = tk.Frame(select_frame)
            search_container.pack(fill="x", pady=5)
            
            tk.Label(search_container, text="Search:", font=("Arial", 9)).pack(side="left", padx=5)
            search_entry = tk.Entry(search_container, textvariable=search_var, width=50)
            search_entry.pack(side="left", fill="x", expand=True)
            
            # Scrollable radio buttons
            list_container = tk.Frame(select_frame)
            list_container.pack(fill="both", expand=True)
            
            canvas = tk.Canvas(list_container)
            scrollbar = tk.Scrollbar(list_container, orient="vertical", command=canvas.yview)
            scrollable_frame = tk.Frame(canvas)
            
            scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            
            selected = tk.StringVar()
            
            def refresh_options(*_):
                for w in scrollable_frame.winfo_children():
                    w.destroy()
                
                term = search_var.get().lower()
                match_count = 0
                
                for display, output in options:
                    if term and term not in display.lower():
                        continue
                    
                    match_count += 1
                    rb = tk.Radiobutton(
                        scrollable_frame,
                        text=display,
                        variable=selected,
                        value=output,
                        wraplength=650,
                        anchor="w",
                        justify="left",
                        font=("Arial", 9)
                    )
                    rb.pack(anchor="w", padx=5, pady=2, fill="x")
                
                if match_count == 0:
                    tk.Label(
                        scrollable_frame,
                        text="No matches found",
                        fg="gray",
                        font=("Arial", 10, "italic")
                    ).pack(pady=20)
            
            search_var.trace_add("write", refresh_options)
            refresh_options()
            
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
            # "This is the last paragraph" checkbox
            checkbox_frame = tk.Frame(dialog, bg="#fff3cd", relief="solid", bd=2)
            checkbox_frame.pack(fill="x", padx=10, pady=10)
            
            is_last = tk.BooleanVar(value=False)
            
            check_container = tk.Frame(checkbox_frame, bg="#fff3cd")
            check_container.pack(pady=10)
            
            tk.Checkbutton(
                check_container,
                text="✓ This is the last paragraph",
                variable=is_last,
                font=("Arial", 11, "bold"),
                bg="#fff3cd",
                activebackground="#fff3cd"
            ).pack(side="left", padx=10)
            
            tk.Label(
                check_container,
                text="(Check this box before clicking 'Add Item' to finish the list)",
                font=("Arial", 8, "italic"),
                bg="#fff3cd",
                fg="#856404"
            ).pack(side="left")
            
            # Bottom buttons
            button_frame = tk.Frame(dialog)
            button_frame.pack(side="bottom", fill="x", pady=15, padx=10)
            
            result = {"value": None, "is_last": False, "cancelled": False}
            
            def on_add_item():
                """Add the selected item and check if this is the last one"""
                val = selected.get()
                if not val:
                    messagebox.showwarning(
                        "No Selection",
                        "Please select an item from the list.",
                        parent=dialog
                    )
                    return
                
                result["value"] = val
                result["is_last"] = is_last.get()
                dialog.destroy()
            
            def on_cancel():
                """Cancel the entire process"""
                if collected_items:
                    confirm = messagebox.askyesno(
                        "Cancel",
                        f"You have {len(collected_items)} item(s) already added.\n\n"
                        "Are you sure you want to cancel and discard all items?",
                        parent=dialog
                    )
                    if not confirm:
                        return
                
                result["cancelled"] = True
                dialog.destroy()
            
            tk.Button(
                button_frame,
                text="Add Item",
                command=on_add_item,
                width=20,
                bg="#4CAF50",
                fg="white",
                font=("Arial", 11, "bold")
            ).pack(side="left", padx=10)
            
            tk.Button(
                button_frame,
                text="Cancel All",
                command=on_cancel,
                width=15,
                bg="#f44336",
                fg="white",
                font=("Arial", 10)
            ).pack(side="right", padx=10)
            
            # Wait for user action
            dialog.wait_window()
            
            # Check if cancelled
            if result["cancelled"]:
                return None
            
            # Check if a value was selected
            if result["value"]:
                collected_items.append(result["value"])
                entry_num += 1
                
                # If this was marked as the last item, STOP the loop
                if result["is_last"]:
                    break
            else:
                # No value selected but dialog closed - treat as cancel
                if not collected_items:
                    return None
                break
        
        # Format as numbered list
        if not collected_items:
            return None
        
        formatted_list = "\n".join([f"{i+1}. {item}" for i, item in enumerate(collected_items)])
        
        return {
            "value": formatted_list,
            "is_single_use": False,
            "use_numbered_list": True,
            "selected_items": collected_items
        }
    
    # ========================================================================
    # TRUE = SINGLE-USE VARIABLE (One selection only)
    # ========================================================================
    else:
        dialog = tk.Toplevel(parent)
        dialog.title(f"Select value for {var_name}")
        dialog.geometry("700x600")
        dialog.grab_set()
        
        # Header
        tk.Label(
            dialog,
            text=f"Select value for: {var_name}",
            font=("Arial", 12, "bold")
        ).pack(pady=10)
        
        tk.Label(
            dialog,
            text="(Single-use variable - will only appear once in document)",
            fg="green",
            font=("Arial", 9)
        ).pack(pady=5)
        
        # Search box
        search_var = tk.StringVar()
        search_frame = tk.Frame(dialog)
        search_frame.pack(fill="x", padx=10, pady=5)
        tk.Label(search_frame, text="Search:").pack(side="left", padx=5)
        tk.Entry(search_frame, textvariable=search_var, width=50).pack(side="left", fill="x", expand=True, padx=5)
        
        # Scrollable options
        canvas = tk.Canvas(dialog)
        scrollbar = tk.Scrollbar(dialog, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        selected = tk.StringVar()
        
        def refresh_options(*_):
            for w in scrollable_frame.winfo_children():
                w.destroy()
            
            term = search_var.get().lower()
            for display, output in options:
                if term and term not in display.lower():
                    continue
                
                rb = tk.Radiobutton(
                    scrollable_frame,
                    text=display,
                    variable=selected,
                    value=output,
                    wraplength=620,
                    anchor="w",
                    justify="left"
                )
                rb.pack(anchor="w", padx=20, pady=2, fill="x")
        
        search_var.trace_add("write", refresh_options)
        refresh_options()
        
        canvas.pack(side="left", fill="both", expand=True, padx=10)
        scrollbar.pack(side="right", fill="y")
        
        result = {"value": None}
        
        def on_submit():
            if not selected.get():
                messagebox.showwarning("No Selection", "Please select an option.", parent=dialog)
                return
            result["value"] = selected.get()
            dialog.destroy()
        
        button_frame = tk.Frame(dialog)
        button_frame.pack(side="bottom", fill="x", pady=10, padx=10)
        tk.Button(button_frame, text="Submit", command=on_submit, width=15, bg="#4CAF50", fg="white").pack()
        
        dialog.wait_window()
        
        return {
            "value": result.get("value"),
            "is_single_use": is_single_use,
            "use_numbered_list": use_numbered_list
        }



def replace_dynamic_variables_in_document(doc_path, replacements):
    """
    Replaces <<variable>> and <<variable_modifier>> in a Word document.
    Preserves formatting. Handles numbered list format for FALSE variables.
    """
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document(doc_path)
    
    def replace_in_paragraph(paragraph, replacements):
        """Replace variables in a paragraph while preserving formatting"""
        for var_name, var_data in replacements.items():
            value = var_data["value"]
            modifier = var_data.get("modifier")
            use_numbered_list = var_data.get("use_numbered_list", False)
            
            # Apply modifiers
            if modifier == "upper":
                value = value.upper()
            elif modifier == "lower":
                value = value.lower()
            elif modifier == "title":
                value = value.title()
            
            # Create pattern for this variable
            if modifier:
                pattern = f"<<{var_name}_{modifier}>>"
            else:
                pattern = f"<<{var_name}>>"
            
            # Replace in paragraph
            if pattern in paragraph.text:
                if use_numbered_list and "\n" in value:
                    # This is a numbered list - need special handling
                    # Replace the pattern with first item, then add remaining as new paragraphs
                    items = value.split("\n")
                    
                    # Replace pattern with first item
                    for run in paragraph.runs:
                        if pattern in run.text:
                            run.text = run.text.replace(pattern, items[0])
                    
                    # Set paragraph formatting for numbered list
                    paragraph.paragraph_format.left_indent = Inches(0.5)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add remaining items as new paragraphs after this one
                    if len(items) > 1:
                        # Insert new paragraphs after current one
                        parent = paragraph._element.getparent()
                        for item in items[1:]:
                            new_p = parent.makeelement('w:p', nsmap=paragraph._element.nsmap)
                            parent.insert(parent.index(paragraph._element) + 1, new_p)
                            new_paragraph = type(paragraph)(new_p, paragraph._parent)
                            new_paragraph.add_run(item)
                            new_paragraph.paragraph_format.left_indent = Inches(0.5)
                            new_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    # Simple replacement - preserves most formatting
                    for run in paragraph.runs:
                        if pattern in run.text:
                            run.text = run.text.replace(pattern, value)
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
    
    # Replace in headers/footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, replacements)
    
    doc.save(doc_path)


# =============================================================================
# VARIABLE TYPE 5: OPPOSING COUNSEL VARIABLES (from DB with (()) delimiters)
# =============================================================================

def extract_opposing_counsel_variables(template_path):
    """
    Extracts opposing counsel variables marked with ((variable)) from a template.
    Returns a set of variable names.
    """
    doc = Document(template_path)
    counsel_vars = set()
    
    # Pattern to match ((variable)) - must have letters, not just parentheses
    # Changed to require at least one letter to avoid matching standalone )
    pattern = r'\(\(([a-zA-Z_][a-zA-Z0-9_]*)\)\)'
    
    for paragraph in doc.paragraphs:
        # Skip paragraphs that are just parentheses
        if paragraph.text.strip() and not re.match(r'^\)+$', paragraph.text.strip()):
            matches = re.findall(pattern, paragraph.text)
            for var_name in matches:
                counsel_vars.add(var_name)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Skip paragraphs that are just parentheses
                    if paragraph.text.strip() and not re.match(r'^\)+$', paragraph.text.strip()):
                        matches = re.findall(pattern, paragraph.text)
                        for var_name in matches:
                            counsel_vars.add(var_name)
    
    return counsel_vars


def select_opposing_counsel_by_id(parent):
    """Let user select an opposing counsel from the database - returns ID"""
    from modules.db import list_opposing_counsel
    
    attorneys = list_opposing_counsel()
    if not attorneys:
        messagebox.showinfo("No Attorneys", "No opposing counsel found in database.\n\nPlease add attorneys first.", parent=parent)
        return None
    
    dialog = tk.Toplevel(parent)
    dialog.title("Select Opposing Counsel")
    dialog.geometry("600x500")
    dialog.grab_set()
    
    tk.Label(dialog, text="Select opposing counsel for this document:", font=("Arial", 12, "bold")).pack(pady=10)
    
    # Search bar at top
    search_frame = tk.Frame(dialog)
    search_frame.pack(fill="x", padx=10, pady=5)
    tk.Label(search_frame, text="Search:").pack(side="left", padx=5)
    search_var = tk.StringVar()
    tk.Entry(search_frame, textvariable=search_var, width=50).pack(side="left", fill="x", expand=True, padx=5)
    
    # Scrollable list in middle
    list_container = tk.Frame(dialog)
    list_container.pack(fill="both", expand=True, padx=10, pady=5)
    
    canvas = tk.Canvas(list_container)
    scrollbar = tk.Scrollbar(list_container, orient="vertical", command=canvas.yview)
    frame = tk.Frame(canvas)
    
    frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)
    
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")
    
    selected = tk.IntVar(value=-1)
    widgets = {}
    
    def refresh_list(*_):
        for widget in frame.winfo_children():
            widget.destroy()
        
        term = search_var.get().lower()
        match_count = 0
        
        for row in attorneys:
            counsel_id = row[0]
            first_name = row[1] or ""
            last_name = row[2] or ""
            firm = row[11] or ""
            
            label = f"{first_name} {last_name}".strip()
            if firm:
                label += f" ({firm})"
            
            if term and term not in label.lower():
                continue
            
            match_count += 1
            
            rb = tk.Radiobutton(
                frame,
                text=label,
                variable=selected,
                value=counsel_id,
                anchor="w",
                wraplength=550
            )
            rb.pack(anchor="w", padx=10, pady=2, fill="x")
        
        if match_count == 0:
            tk.Label(frame, text="No matching attorneys", fg="gray").pack(pady=20)
    
    search_var.trace_add("write", refresh_list)
    refresh_list()
    
    # Submit button at bottom
    button_frame = tk.Frame(dialog)
    button_frame.pack(side="bottom", fill="x", pady=10)
    
    result = {"id": None}
    
    def on_submit():
        if selected.get() == -1:
            messagebox.showwarning("No Selection", "Please select an attorney.", parent=dialog)
            return
        result["id"] = selected.get()
        dialog.destroy()
    
    tk.Button(button_frame, text="Select Attorney", command=on_submit, width=20, bg="#4CAF50", fg="white").pack()
    
    dialog.wait_window()
    
    return result["id"]


def replace_opposing_counsel_variables(doc_path, counsel_data, counsel_id, parent_window=None):
    """
    Replace ((variable)) with opposing counsel data.
    Prompts for missing fields and saves them to the database.
    Works exactly like {{}} variables - if not found, prompt and save.
    """
    from modules.db import update_opposing_counsel, get_opposing_counsel, DB_PATH
    import sqlite3
    from tkinter import simpledialog
    from docx import Document
    import re
    
    doc = Document(doc_path)
    
    # Collect all ((variables)) in document
    all_vars_in_doc = set()
    pattern = r'\(\(([a-zA-Z_][a-zA-Z0-9_]*)\)\)'
    
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        for var_name in matches:
            all_vars_in_doc.add(var_name.lower())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    for var_name in matches:
                        all_vars_in_doc.add(var_name.lower())
    
    # Build lowercase counsel_data
    counsel_lower = {k.lower(): v for k, v in counsel_data.items()}
    
    # Track if we need to update the database
    updates_needed = False
    new_fields = {}
    
    # Check for missing variables and prompt
    for var_name in all_vars_in_doc:
        if var_name not in counsel_lower or not counsel_lower[var_name]:
            if parent_window:
                response = messagebox.askyesnocancel(
                    "Missing Attorney Field",
                    f"The attorney field '(({var_name}))' is not in the database.\n\n"
                    f"Would you like to:\n"
                    f"  • YES: Enter a value and save to attorney database\n"
                    f"  • NO: Leave as (({var_name})) in document\n"
                    f"  • CANCEL: Stop document generation",
                    parent=parent_window
                )
                
                if response is None:  # Cancel
                    raise Exception("User cancelled due to missing attorney variable")
                elif response:  # Yes - prompt and save
                    value = simpledialog.askstring(
                        "Enter Value",
                        f"Enter value for (({var_name})):",
                        parent=parent_window
                    )
                    if value:
                        counsel_lower[var_name] = value
                        new_fields[var_name] = value
                        updates_needed = True
                    else:
                        counsel_lower[var_name] = ""  # User entered empty - that's ok
                else:  # No - leave as placeholder, don't stop
                    counsel_lower[var_name] = f"(({var_name}))"  # Leave unchanged
    
    # If we collected new fields, save them to the database
    if updates_needed and new_fields:
        # Get current attorney data
        attorney_row = get_opposing_counsel(counsel_id)
        if attorney_row:
            # Map variable names back to database columns
            field_mapping = {
                "plaintiffattorneyfirstname": "first_name",
                "plaintiffattorneylastname": "last_name",
                "plaintiffattorneyemail": "email",
                "plaintiffattorneyeserviceemail": "service_email",
                "plaintifffirmaddress": "address_street",
                "plaintifffirmcity": "address_city",
                "plaintifffirmst": "address_state",
                "plaintifffirmzip": "address_zip",
                "plaintiffbusphone": "phone",
                "plaintifffaxphone": "fax",
                "plaintifffirmname": "firm_name",
                "plaintiffbarnumber": "bar_number",
                "plaintiffnotes": "notes",
            }
            
            # Build update values
            update_data = {
                "first_name": attorney_row[1],
                "last_name": attorney_row[2],
                "email": attorney_row[3],
                "service_email": attorney_row[4],
                "address_street": attorney_row[5],
                "address_city": attorney_row[6],
                "address_state": attorney_row[7],
                "address_zip": attorney_row[8],
                "phone": attorney_row[9],
                "fax": attorney_row[10],
                "firm_name": attorney_row[11],
                "bar_number": attorney_row[12],
                "notes": attorney_row[13],
            }
            
            # Apply new field values
            for var_name, value in new_fields.items():
                if var_name in field_mapping:
                    db_field = field_mapping[var_name]
                    update_data[db_field] = value
                else:
                    # Custom field - add to notes
                    if update_data["notes"]:
                        update_data["notes"] += f"\n{var_name}: {value}"
                    else:
                        update_data["notes"] = f"{var_name}: {value}"
            
            # Save to database
            update_opposing_counsel(
                counsel_id,
                update_data["first_name"],
                update_data["last_name"],
                update_data["email"],
                update_data["service_email"],
                update_data["address_street"],
                update_data["address_city"],
                update_data["address_state"],
                update_data["address_zip"],
                update_data["phone"],
                update_data["fax"],
                update_data["firm_name"],
                update_data["bar_number"],
                update_data["notes"]
            )
            
            messagebox.showinfo(
                "Attorney Updated",
                f"Saved {len(new_fields)} new field(s) to attorney database.",
                parent=parent_window
            )
    
    # Now do replacement
    def replace_in_paragraph(paragraph):
        matches = re.findall(pattern, paragraph.text)
        
        for var_name in matches:
            var_name_lower = var_name.lower()
            
            if var_name_lower in counsel_lower and counsel_lower[var_name_lower]:
                value = counsel_lower[var_name_lower]
                original_pattern = f"(({var_name}))"
                
                for run in paragraph.runs:
                    if original_pattern in run.text:
                        run.text = run.text.replace(original_pattern, str(value))
    
    # Replace in all document sections
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph)
    
    doc.save(doc_path)


# =============================================================================
# VARIABLE TYPE 6: DOCUMENT-SPECIFIC VARIABLES (prompted each time with {@@})
# =============================================================================

def extract_document_specific_variables(template_path):
    """
    Extracts document-specific variables marked with {@variable@} from a template.
    These are prompted every time during document generation.
    Returns a set of variable names.
    """
    doc = Document(template_path)
    doc_vars = set()
    
    # Pattern to match {@variable@}
    pattern = r'\{@([a-zA-Z_][a-zA-Z0-9_]*)@\}'
    
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        for var_name in matches:
            doc_vars.add(var_name)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    for var_name in matches:
                        doc_vars.add(var_name)
    
    return doc_vars


def prompt_document_specific_variable(parent, var_name):
    """Prompt for a document-specific variable"""
    dialog = tk.Toplevel(parent)
    dialog.title(f"Document-Specific: {var_name}")
    dialog.geometry("500x200")
    dialog.grab_set()
    
    tk.Label(dialog, text=f"Enter value for: {var_name}", font=("Arial", 12, "bold")).pack(pady=10)
    tk.Label(dialog, text="(This will be prompted for each document)", fg="blue").pack(pady=5)
    
    value_var = tk.StringVar()
    tk.Label(dialog, text="Value:").pack(pady=5)
    entry = tk.Entry(dialog, textvariable=value_var, width=50)
    entry.pack(pady=5)
    entry.focus()
    
    result = {"value": None}
    
    def on_submit():
        val = value_var.get().strip()
        if not val:
            messagebox.showwarning("Empty Value", "Please enter a value.", parent=dialog)
            return
        result["value"] = val
        dialog.destroy()
    
    tk.Button(dialog, text="Submit", command=on_submit, width=15).pack(pady=10)
    entry.bind("<Return>", lambda e: on_submit())
    
    dialog.wait_window()
    
    return result["value"] or ""


def replace_document_specific_variables(doc_path, doc_vars_data):
    """Replace {@variable@} with document-specific data"""
    doc = Document(doc_path)
    
    def replace_in_paragraph(paragraph):
        for var_name, value in doc_vars_data.items():
            pattern = f"{{@{var_name}@}}"
            if pattern in paragraph.text:
                for run in paragraph.runs:
                    if pattern in run.text:
                        run.text = run.text.replace(pattern, str(value))
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    
    # Replace in headers/footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph)
    
    doc.save(doc_path)








# =============================================================================
# CLIENT SELECTION HELPERS
# =============================================================================

def build_client_label(client_id):
    """Build a descriptive label for a client"""
    vars_ = get_variables("client", client_id)
    parts = []
    
    # Add matter ID if exists
    if vars_.get("matterid"):
        parts.append(f"Matter: {vars_.get('matterid')}")
    
    # Add name if exists
    fname = vars_.get("firstname", "")
    lname = vars_.get("lastname", "")
    if fname or lname:
        parts.append(f"{fname} {lname}".strip())
    
    label = f"ID {client_id}"
    if parts:
        label += " | " + " | ".join(parts)
    
    return label


def select_client(clients):
    """Display client selection dialog"""
    selected_id = tk.IntVar(value=-1)
    
    popup = tk.Toplevel()
    popup.title("Select Client")
    popup.geometry("800x600")
    popup.grab_set()
    
    tk.Label(popup, text="Select a client:", font=("Arial", 12, "bold")).pack(pady=10)
    
    # Search functionality
    search_frame = tk.Frame(popup)
    search_frame.pack(fill="x", padx=10, pady=5)
    search_var = tk.StringVar()
    tk.Label(search_frame, text="Search:").pack(side="left", padx=5)
    tk.Entry(search_frame, textvariable=search_var, width=60).pack(side="left", fill="x", expand=True, padx=5)
    
    # Scrollable frame for clients
    canvas = tk.Canvas(popup)
    scrollbar = tk.Scrollbar(popup, orient="vertical", command=canvas.yview)
    frame = tk.Frame(canvas)
    
    frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)
    
    def refresh(*_):
        """Refresh client list based on search"""
        for w in frame.winfo_children():
            w.destroy()
        
        term = search_var.get().lower()
        match_count = 0
        
        for cid, *_ in clients:
            label = build_client_label(cid)
            if term in label.lower():
                tk.Radiobutton(
                    frame,
                    text=label,
                    variable=selected_id,
                    value=cid,
                    wraplength=750,
                    anchor="w",
                    justify="left"
                ).pack(anchor="w", padx=10, pady=2, fill="x")
                match_count += 1
        
        if match_count == 0:
            tk.Label(frame, text="No matching clients found", fg="gray").pack(padx=10, pady=20)
    
    search_var.trace_add("write", refresh)
    refresh()
    
    canvas.pack(side="left", fill="both", expand=True, padx=10)
    scrollbar.pack(side="right", fill="y")
    
    def submit():
        if selected_id.get() == -1:
            messagebox.showwarning("No Selection", "Please select a client.")
            return
        popup.destroy()
    
    tk.Button(popup, text="Select Client", command=submit, width=15).pack(pady=10)
    
    popup.wait_window()
    return selected_id.get() if selected_id.get() != -1 else None


def select_all_clients_option(clients):
    """Ask if user wants to generate for all clients or select one"""
    dialog = tk.Toplevel()
    dialog.title("Client Selection")
    dialog.geometry("400x200")
    dialog.grab_set()
    
    tk.Label(dialog, text="Generate documents for:", font=("Arial", 12, "bold")).pack(pady=20)
    
    choice = tk.StringVar(value="single")
    
    tk.Radiobutton(dialog, text="Single client", variable=choice, value="single").pack(pady=5)
    tk.Radiobutton(dialog, text="All clients", variable=choice, value="all").pack(pady=5)
    
    result = {"choice": None}
    
    def submit():
        result["choice"] = choice.get()
        dialog.destroy()
    
    tk.Button(dialog, text="Continue", command=submit, width=15).pack(pady=20)
    
    dialog.wait_window()
    
    if result["choice"] == "all":
        return [c[0] for c in clients]
    else:
        selected = select_client(clients)
        return [selected] if selected else []


# =============================================================================
# TEMPLATE SELECTION
# =============================================================================

def select_templates(templates_dir="templates"):
    """Display template selection dialog with checkboxes"""
    templates = sorted(Path(templates_dir).glob("*.docx"))
    
    if not templates:
        messagebox.showinfo("No Templates", "No .docx templates found in templates folder.")
        return []
    
    window = tk.Toplevel()
    window.title("Select Templates")
    window.geometry("700x600")
    window.grab_set()
    
    tk.Label(window, text="Select templates to generate:", font=("Arial", 12, "bold")).pack(pady=10, anchor="w", padx=10)
    
    # Search box
    search_var = tk.StringVar()
    search_frame = tk.Frame(window)
    search_frame.pack(fill="x", padx=10, pady=5)
    tk.Label(search_frame, text="Search:").pack(side="left", padx=5)
    tk.Entry(search_frame, textvariable=search_var, width=50).pack(side="left", padx=5)
    
    # Container for canvas and scrollbar
    canvas_container = tk.Frame(window)
    canvas_container.pack(fill="both", expand=True, padx=10, pady=5)
    
    # Create scrollable frame INSIDE the container
    canvas = tk.Canvas(canvas_container)
    scrollbar = tk.Scrollbar(canvas_container, orient="vertical", command=canvas.yview)
    frame = tk.Frame(canvas)
    
    frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)
    
    selections = {}
    checkbuttons = {}
    
    def refresh_list(*_):
        search_term = search_var.get().lower()
        for t, cb in checkbuttons.items():
            if search_term in t.name.lower():
                cb.pack(anchor="w", padx=5, pady=2, fill="x")
            else:
                cb.pack_forget()
    
    for t in templates:
        var = tk.BooleanVar()
        selections[t] = var
        cb = tk.Checkbutton(frame, text=t.name, variable=var, anchor="w", justify="left")
        checkbuttons[t] = cb
        cb.pack(anchor="w", padx=5, pady=2, fill="x")
    
    search_var.trace_add("write", refresh_list)
    
    # Pack canvas and scrollbar side-by-side WITHIN the container
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")
    
    result = {"templates": []}
    
    def on_submit():
        selected = [t for t, v in selections.items() if v.get()]
        if not selected:
            messagebox.showwarning("No Selection", "Please select at least one template.", parent=window)
            return
        result["templates"] = selected
        window.destroy()
    
    def select_all():
        for var in selections.values():
            var.set(True)
    
    def deselect_all():
        for var in selections.values():
            var.set(False)
    
    # Bottom button frame
    bottom_frame = tk.Frame(window)
    bottom_frame.pack(fill="x", side="bottom", padx=10, pady=10)
    
    # Buttons left-aligned
    tk.Button(bottom_frame, text="Select All", command=select_all, width=12).pack(side="left", padx=5)
    tk.Button(bottom_frame, text="Deselect All", command=deselect_all, width=12).pack(side="left", padx=5)
    tk.Button(bottom_frame, text="Generate", command=on_submit, width=12).pack(side="left", padx=5)
    
    window.wait_window()
    
    return result["templates"]


# =============================================================================
# MAIN DOCUMENT GENERATION LOGIC
# =============================================================================
def generate_document_from_template(template_path, client_id, parent_window=None):
    """
    Generate a document from a template for a specific client.
    Handles all variable types IN PRIORITY ORDER.
    """
    output_dir = Path("output_documents")
    output_dir.mkdir(exist_ok=True)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_file = output_dir / f"{template_path.stem}_client{client_id}_{timestamp}.docx"
    
    import shutil
    shutil.copy(template_path, output_file)
    working_copy = output_dir / f"_temp_{template_path.stem}_{timestamp}.docx"
    shutil.copy(template_path, working_copy)
    
    all_client_vars = get_variables("client", client_id)
    
    # ===================================================================
    # STEP 1: Handle <<angle bracket>> dynamic variables FIRST
    # This ensures venue/Jurisdiction are captured before {{ }} processing
    # ===================================================================
    dynamic_vars = extract_dynamic_variables_from_template(working_copy)
    
    if dynamic_vars:
        replacements = {}
        prompted_base_vars = set()
        
        for var_name, modifier in dynamic_vars:
            full_key = f"{var_name}_{modifier}" if modifier else var_name
            
            if full_key in replacements:
                continue
            
            if var_name in prompted_base_vars:
                if var_name in replacements:
                    base_value = replacements[var_name]["value"]
                    modified_value = base_value
                    if modifier == "upper":
                        modified_value = base_value.upper()
                    elif modifier == "lower":
                        modified_value = base_value.lower()
                    elif modifier == "title":
                        modified_value = base_value.title()
                    
                    replacements[full_key] = {
                        "value": modified_value,
                        "modifier": modifier,
                        "is_single_use": replacements[var_name].get("is_single_use", False),
                        "use_numbered_list": replacements[var_name].get("use_numbered_list", False)
                    }
                    
                    # ONLY store BASE variable in all_client_vars (NOT the modified one)
                    # all_client_vars[full_key] = modified_value  # REMOVE THIS LINE
                continue
            
            prompted_base_vars.add(var_name)
            result = prompt_dynamic_variable_from_excel(parent_window, var_name, client_id)
            
            if result and result["value"]:
                replacements[var_name] = {
                    "value": result["value"],
                    "modifier": None,
                    "is_single_use": result.get("is_single_use", False),
                    "use_numbered_list": result.get("use_numbered_list", False)
                }
                
                replacements[full_key] = {
                    "value": result["value"],
                    "modifier": modifier,
                    "is_single_use": result.get("is_single_use", False),
                    "use_numbered_list": result.get("use_numbered_list", False)
                }
                
                # ONLY store BASE variable
                set_variable("client", client_id, var_name, result["value"])
                all_client_vars[var_name] = result["value"]
                # all_client_vars[full_key] = result["value"]  # REMOVE THIS LINE
        
        # Replace dynamic variables in WORKING COPY
        replace_dynamic_variables_in_document(working_copy, replacements)
    
    # ===================================================================
    # STEP 2: Handle {{double brace}} variables with docxtpl
    # Now all_client_vars contains the <<>> values too
    # ===================================================================
    try:
        tpl = DocxTemplate(working_copy)
        
        try:
            raw_vars = tpl.get_undeclared_template_variables()
        except Exception as e:
            # ... (keep your existing filter conversion code)
            error_msg = str(e)
            if "No filter named" in error_msg or "filter" in error_msg.lower():
                # ... existing conversion logic ...
                pass
        
        context = {}
        context.update(get_system_date_context())
        
        for placeholder in raw_vars:
            if placeholder in context:
                continue
            
            # Detect suffixes/modifiers
            var_name = placeholder
            modifiers = []
            
            if var_name.endswith("_combo"):
                modifiers.append("combo")
                var_name = var_name.rsplit("_combo", 1)[0]
            
            if var_name.endswith("_derived"):
                modifiers.append("derived")
                var_name = var_name.rsplit("_derived", 1)[0]
            
            if var_name.endswith("_upper"):
                modifiers.append("upper")
                var_name = var_name.rsplit("_upper", 1)[0]
            
            if var_name.endswith("_lower"):
                modifiers.append("lower")
                var_name = var_name.rsplit("_lower", 1)[0]
            
            if var_name.endswith("_title"):
                modifiers.append("title")
                var_name = var_name.rsplit("_title", 1)[0]
            
            value = all_client_vars.get(var_name)
            
            if value is None or value == "":
                # FIRST: Check if it's a combo variable
                if "combo" in modifiers:
                    concats = {c["var_name"]: c for c in list_all_concats()}
                    if var_name in concats:
                        value = handle_concatenated_variable(parent_window, var_name, client_id, all_client_vars)
                    else:
                        # Not defined yet - open editor
                        value = handle_concatenated_variable(parent_window, var_name, client_id, all_client_vars)
                # SECOND: Check if it's defined as a concat variable (without tag)
                elif var_name in {c["var_name"]: c for c in list_all_concats()}:
                    value = handle_concatenated_variable(parent_window, var_name, client_id, all_client_vars)
                # THIRD: Check if it's a derived/grammatical variable
                elif "derived" in modifiers:
                    value = handle_derived_variable(parent_window, var_name, client_id, all_client_vars)
                # LAST: Prompt user for value
                else:
                    value = prompt_for_variable(parent_window, var_name, client_id, all_client_vars)
                    all_client_vars[var_name] = value
            
            if value:
                if "upper" in modifiers:
                    value = str(value).upper()
                elif "lower" in modifiers:
                    value = str(value).lower()
                elif "title" in modifiers:
                    value = str(value).title()
            
            context[placeholder] = value if value else ""
        
        tpl.render(context)
        tpl.save(output_file)
        
        if working_copy.exists():
            working_copy.unlink()
        
    except Exception as e:
        if working_copy.exists():
            working_copy.unlink()
        messagebox.showerror("Document Generation Error", f"Failed: {e}", parent=parent_window)
        return None
    
    # Step 3: Handle ((double parenthesis)) opposing counsel variables
    counsel_vars = extract_opposing_counsel_variables(output_file)
    
    if counsel_vars:
        from modules.db import get_opposing_counsel_variables, DB_PATH
        import sqlite3
        
        # Read opposing_counsel_id DIRECTLY from clients table
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT opposing_counsel_id FROM clients WHERE id = ?", (client_id,))
        row = c.fetchone()
        conn.close()
        
        assigned_counsel_id = row[0] if row and row[0] else None
        
        if assigned_counsel_id:
            try:
                counsel_data = get_opposing_counsel_variables(int(assigned_counsel_id))
                replace_opposing_counsel_variables(output_file, counsel_data, int(assigned_counsel_id), parent_window)  # PASS counsel_id
            except Exception as e:
                messagebox.showwarning("Attorney Error", f"Could not load attorney.\n\n{e}", parent=parent_window)
                counsel_id = select_opposing_counsel_by_id(parent_window)
                if counsel_id:
                    counsel_data = get_opposing_counsel_variables(counsel_id)
                    replace_opposing_counsel_variables(output_file, counsel_data, counsel_id, parent_window)  # PASS counsel_id
        else:
            messagebox.showinfo("Attorney Required", "Please select an attorney.", parent=parent_window)
            counsel_id = select_opposing_counsel_by_id(parent_window)
            if counsel_id:
                counsel_data = get_opposing_counsel_variables(counsel_id)
                replace_opposing_counsel_variables(output_file, counsel_data, counsel_id, parent_window)  # PASS counsel_id
                # Save assignment
                conn = sqlite3.connect(DB_PATH)
                c = conn.cursor()
                c.execute("UPDATE clients SET opposing_counsel_id = ? WHERE id = ?", (counsel_id, client_id))
                conn.commit()
                conn.close()



    # Step 3.5: Handle (@grammar@) variables
    grammar_vars = extract_grammar_variables(output_file)
    
    if grammar_vars:
        grammar_settings = prompt_grammar_settings(parent_window)
        if grammar_settings["count"]:
            replace_grammar_variables(output_file, grammar_settings)




    # Step 4: Document-specific variables
    doc_specific_vars = extract_document_specific_variables(output_file)
    
    if doc_specific_vars:
        doc_vars_data = {}
        for var_name in doc_specific_vars:
            value = prompt_document_specific_variable(parent_window, var_name)
            doc_vars_data[var_name] = value
        replace_document_specific_variables(output_file, doc_vars_data)
    


        # Step 5: Handle [[bracket]] variables from Excel dynamic content
    bracket_vars = extract_bracket_variables(output_file)
    
    if bracket_vars:
        replace_bracket_variables(output_file, client_id)
    
    return str(output_file)







def generate_documents(client_id=None):
    """
    Main orchestrator function called from GUI.
    Handles client selection, template selection, and document generation.
    """
    # Get all clients
    clients = list_clients()
    if not clients:
        messagebox.showinfo("No Clients", "No clients found in database.")
        return
    
    # Step 1: Select client(s)
    if client_id is None:
        client_ids = select_all_clients_option(clients)
        if not client_ids:
            return
    else:
        client_ids = [client_id]
    
    # Step 2: Select templates
    selected_templates = select_templates()
    if not selected_templates:
        return
    
    # Step 3: Generate documents
    generated_files = []
    total_docs = len(client_ids) * len(selected_templates)
    
    # Progress tracking
    progress_window = tk.Toplevel()
    progress_window.title("Generating Documents")
    progress_window.geometry("400x150")
    progress_window.grab_set()
    
    tk.Label(progress_window, text="Generating documents...", font=("Arial", 12)).pack(pady=10)
    progress_var = tk.StringVar(value="0 of " + str(total_docs))
    tk.Label(progress_window, textvariable=progress_var).pack(pady=5)
    
    progress_bar = ttk.Progressbar(progress_window, length=300, mode='determinate')
    progress_bar.pack(pady=10)
    progress_bar['maximum'] = total_docs
    
    count = 0
    
    for client_id in client_ids:
        for template in selected_templates:
            progress_var.set(f"{count + 1} of {total_docs}")
            progress_bar['value'] = count + 1
            progress_window.update()
            
            try:
                output_file = generate_document_from_template(template, client_id, progress_window)
                if output_file:
                    generated_files.append(output_file)
            except Exception as e:
                messagebox.showerror(
                    "Generation Error",
                    f"Error generating {template.name} for client {client_id}:\n{e}",
                    parent=progress_window
                )
            
            count += 1
    
    progress_window.destroy()
    
    # Show summary
    if generated_files:
        messagebox.showinfo(
            "Success",
            f"Successfully generated {len(generated_files)} document(s).\n\n"
            f"Output folder: output_documents/"
        )
        
        # Ask if user wants to open output folder
        if messagebox.askyesno("Open Folder", "Would you like to open the output folder?"):
            import subprocess
            import sys
            output_dir = Path("output_documents").resolve()
            
            if sys.platform == "win32":
                subprocess.Popen(["explorer", str(output_dir)])
            elif sys.platform == "darwin":
                subprocess.Popen(["open", str(output_dir)])
            else:
                subprocess.Popen(["xdg-open", str(output_dir)])
    else:
        messagebox.showwarning("No Documents", "No documents were generated.")


# =============================================================================
# ENTRY POINT (called from main.py)
# =============================================================================

if __name__ == "__main__":
    # For testing
    generate_documents()