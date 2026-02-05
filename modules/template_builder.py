# modules/template_builder.py
"""
Template Builder - Convert existing documents into variable templates.
Allows users to select text and replace it with variable placeholders.
"""

import tkinter as tk
from tkinter import messagebox, filedialog, ttk, simpledialog
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
import re


class TemplateBuilder:
    def __init__(self, parent=None):
        self.window = tk.Toplevel(parent) if parent else tk.Tk()
        self.window.title("Template Builder")
        self.window.geometry("1400x900")
        
        self.doc = None
        self.doc_path = None
        self.current_text_widget = None
        self.replacements = []  # Track all replacements made
        
        self.setup_ui()
    
    def setup_ui(self):
        """Setup the UI layout"""
        # Top toolbar
        toolbar = tk.Frame(self.window, bg="#f0f0f0", height=60)
        toolbar.pack(side="top", fill="x")
        toolbar.pack_propagate(False)
        
        tk.Button(
            toolbar,
            text="📁 Open Document",
            command=self.load_document,
            width=15,
            height=2
        ).pack(side="left", padx=10, pady=10)
        
        tk.Button(
            toolbar,
            text="💾 Save Template",
            command=self.save_template,
            width=15,
            height=2,
            bg="#4CAF50",
            fg="white"
        ).pack(side="left", padx=10, pady=10)
        
        tk.Button(
            toolbar,
            text="↩️ Undo Last",
            command=self.undo_last,
            width=12,
            height=2
        ).pack(side="left", padx=10, pady=10)
        
        self.status_label = tk.Label(
            toolbar,
            text="No document loaded",
            font=("Arial", 10),
            bg="#f0f0f0"
        )
        self.status_label.pack(side="left", padx=20)
        
        # Main content area
        main_frame = tk.Frame(self.window)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Left side: Document preview/editor
        left_frame = tk.Frame(main_frame)
        left_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))
        
        tk.Label(
            left_frame,
            text="Document Preview (select text to replace)",
            font=("Arial", 11, "bold")
        ).pack(anchor="w", pady=5)
        
        # Text widget with scrollbar
        text_container = tk.Frame(left_frame)
        text_container.pack(fill="both", expand=True)
        
        text_scroll = tk.Scrollbar(text_container)
        text_scroll.pack(side="right", fill="y")
        
        self.text_widget = tk.Text(
            text_container,
            wrap="word",
            yscrollcommand=text_scroll.set,
            font=("Arial", 11),
            bg="#ffffff",
            padx=10,
            pady=10
        )
        self.text_widget.pack(side="left", fill="both", expand=True)
        text_scroll.config(command=self.text_widget.yview)
        
        # Bind right-click
        self.text_widget.bind("<Button-3>", self.show_replace_menu)
        
        # Configure tags for highlighting
        self.text_widget.tag_configure("variable", background="#ffeb3b", foreground="#000000")
        self.text_widget.tag_configure("selected", background="#2196F3", foreground="#ffffff")
        
        # Right side: Variable palette and replacements log
        right_frame = tk.Frame(main_frame, width=400)
        right_frame.pack(side="right", fill="both")
        right_frame.pack_propagate(False)
        
        # Variable palette
        tk.Label(
            right_frame,
            text="Quick Variable Palette",
            font=("Arial", 11, "bold")
        ).pack(anchor="w", pady=5)
        
        palette_frame = tk.Frame(right_frame)
        palette_frame.pack(fill="x", pady=5)
        
        # Common variable buttons
        common_vars = [
            ("{{firstname}}", "First Name"),
            ("{{lastname}}", "Last Name"),
            ("{{matterid}}", "Matter ID"),
            ("<<Jurisdiction>>", "Jurisdiction"),
            ("<<venue>>", "Venue"),
            ("((plaintiffattorneyfullname))", "Attorney Name"),
            ("{{currentday}}", "Current Day"),
            ("{{currentmonth}}", "Current Month"),
            ("{{year}}", "Year"),
        ]
        
        for i, (var, label) in enumerate(common_vars):
            btn = tk.Button(
                palette_frame,
                text=label,
                command=lambda v=var: self.replace_selection(v),
                width=18,
                anchor="w"
            )
            btn.grid(row=i // 2, column=i % 2, padx=2, pady=2, sticky="ew")
        
        palette_frame.columnconfigure(0, weight=1)
        palette_frame.columnconfigure(1, weight=1)
        
        # Custom variable entry
        tk.Label(right_frame, text="Custom Variable:", font=("Arial", 10, "bold")).pack(anchor="w", pady=(20, 5))
        
        custom_frame = tk.Frame(right_frame)
        custom_frame.pack(fill="x", pady=5)
        
        self.var_type_var = tk.StringVar(value="{{")
        tk.OptionMenu(custom_frame, self.var_type_var, "{{", "<<", "((" , "{@", "[[", "(@").pack(side="left", padx=2)
        
        self.custom_var_entry = tk.Entry(custom_frame, width=15)
        self.custom_var_entry.pack(side="left", fill="x", expand=True, padx=2)
        
        tk.Button(
            custom_frame,
            text="Insert",
            command=self.insert_custom_variable,
            width=8
        ).pack(side="left", padx=2)
        
        # Replacements log
        tk.Label(
            right_frame,
            text="Replacements Made:",
            font=("Arial", 11, "bold")
        ).pack(anchor="w", pady=(20, 5))
        
        log_container = tk.Frame(right_frame)
        log_container.pack(fill="both", expand=True)
        
        log_scroll = tk.Scrollbar(log_container)
        log_scroll.pack(side="right", fill="y")
        
        self.log_text = tk.Text(
            log_container,
            wrap="word",
            yscrollcommand=log_scroll.set,
            font=("Courier", 9),
            bg="#f5f5f5",
            height=15,
            state="disabled"
        )
        self.log_text.pack(side="left", fill="both", expand=True)
        log_scroll.config(command=self.log_text.yview)
    
    def load_document(self):
        """Load a .docx document"""
        filepath = filedialog.askopenfilename(
            title="Select Document to Convert",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not filepath:
            return
        
        try:
            self.doc = Document(filepath)
            self.doc_path = Path(filepath)
            self.replacements = []
            
            # Extract text from document
            full_text = []
            for paragraph in self.doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Display in text widget
            self.text_widget.delete("1.0", "end")
            self.text_widget.insert("1.0", "\n".join(full_text))
            
            self.status_label.config(text=f"Loaded: {self.doc_path.name}")
            
            self.log_message(f"✓ Loaded document: {self.doc_path.name}")
            
        except Exception as e:
            messagebox.showerror("Load Error", f"Failed to load document:\n{e}")
    
    def show_replace_menu(self, event):
        """Show context menu when right-clicking selected text"""
        try:
            selection = self.text_widget.get("sel.first", "sel.last")
            if not selection:
                return
            
            # Create popup menu
            menu = tk.Menu(self.window, tearoff=0)
            
            menu.add_command(
                label=f"Replace '{selection[:30]}...' with variable",
                command=self.show_variable_dialog
            )
            
            menu.add_separator()
            
            # Quick replace options
            menu.add_command(
                label="Client Variable {{...}}",
                command=lambda: self.quick_replace("{{")
            )
            menu.add_command(
                label="Dynamic Variable <<...>>",
                command=lambda: self.quick_replace("<<")
            )
            menu.add_command(
                label="Attorney Variable ((...)))",
                command=lambda: self.quick_replace("((")
            )
            menu.add_command(
                label="Grammar Variable (@...@)",
                command=lambda: self.quick_replace("(@")
            )
            
            menu.post(event.x_root, event.y_root)
            
        except tk.TclError:
            # No selection
            pass
    
    def show_variable_dialog(self):
        """Show dialog to choose variable type and name"""
        try:
            selection = self.text_widget.get("sel.first", "sel.last")
        except tk.TclError:
            messagebox.showwarning("No Selection", "Please select text first.")
            return
        
        dialog = tk.Toplevel(self.window)
        dialog.title("Replace with Variable")
        dialog.geometry("500x350")
        dialog.grab_set()
        
        tk.Label(
            dialog,
            text=f"Replace: \"{selection[:50]}...\"",
            font=("Arial", 11, "bold")
        ).pack(pady=10)
        
        # Variable type
        tk.Label(dialog, text="Variable Type:").pack(anchor="w", padx=20, pady=(10, 0))
        var_type = tk.StringVar(value="client")
        
        type_frame = tk.Frame(dialog)
        type_frame.pack(fill="x", padx=20, pady=5)
        
        tk.Radiobutton(type_frame, text="Client Variable {{...}}", variable=var_type, value="client").pack(anchor="w")
        tk.Radiobutton(type_frame, text="Dynamic Variable <<...>>", variable=var_type, value="dynamic").pack(anchor="w")
        tk.Radiobutton(type_frame, text="Attorney Variable ((...)))", variable=var_type, value="attorney").pack(anchor="w")
        tk.Radiobutton(type_frame, text="Document-Specific {@...@}", variable=var_type, value="docspec").pack(anchor="w")
        tk.Radiobutton(type_frame, text="Excel Dynamic [[...]]", variable=var_type, value="bracket").pack(anchor="w")
        tk.Radiobutton(type_frame, text="Grammar Variable (@...@)", variable=var_type, value="grammar").pack(anchor="w")
        
        # Variable name
        tk.Label(dialog, text="Variable Name:").pack(anchor="w", padx=20, pady=(10, 0))
        
        var_name_entry = tk.Entry(dialog, width=40)
        var_name_entry.pack(padx=20, pady=5)
        var_name_entry.focus()
        
        # Auto-suggest name from selection
        suggested_name = re.sub(r'[^a-zA-Z0-9_]', '', selection.lower().replace(" ", "_"))[:30]
        var_name_entry.insert(0, suggested_name)
        
        result = {"variable": None}
        
        def on_create():
            var_name = var_name_entry.get().strip()
            if not var_name:
                messagebox.showwarning("Empty Name", "Please enter a variable name.", parent=dialog)
                return
            
            # Build variable based on type
            vtype = var_type.get()
            if vtype == "client":
                result["variable"] = f"{{{{{var_name}}}}}"
            elif vtype == "dynamic":
                result["variable"] = f"<<{var_name}>>"
            elif vtype == "attorney":
                result["variable"] = f"(({var_name}))"
            elif vtype == "docspec":
                result["variable"] = f"{{@{var_name}@}}"
            elif vtype == "bracket":
                result["variable"] = f"[[{var_name}]]"
            elif vtype == "grammar":
                result["variable"] = f"(@{var_name}@)"
            
            dialog.destroy()
        
        tk.Button(dialog, text="Create Variable", command=on_create, width=20, bg="#4CAF50", fg="white").pack(pady=20)
        
        dialog.wait_window()
        
        if result["variable"]:
            self.replace_selection(result["variable"])
    
    def quick_replace(self, delimiter):
        """Quick replace with just delimiter, prompting for name"""
        try:
            selection = self.text_widget.get("sel.first", "sel.last")
        except tk.TclError:
            return
        
        var_name = simpledialog.askstring(
            "Variable Name",
            f"Enter variable name for '{selection[:30]}':",
            parent=self.window
        )
        
        if not var_name:
            return
        
        # Build variable based on delimiter
        delim_map = {
            "{{": f"{{{{{var_name}}}}}",
            "<<": f"<<{var_name}>>",
            "((": f"(({var_name}))",
            "{@": f"{{@{var_name}@}}",
            "[[": f"[[{var_name}]]",
            "(@": f"(@{var_name}@)"
        }
        
        variable = delim_map.get(delimiter, f"{{{{{var_name}}}}}")
        self.replace_selection(variable)
    
    def insert_custom_variable(self):
        """Insert custom variable from entry field"""
        var_name = self.custom_var_entry.get().strip()
        if not var_name:
            messagebox.showwarning("Empty Name", "Please enter a variable name.")
            return
        
        delimiter = self.var_type_var.get()
        
        # Build variable
        delim_map = {
            "{{": f"{{{{{var_name}}}}}",
            "<<": f"<<{var_name}>>",
            "((": f"(({var_name}))",
            "{@": f"{{@{var_name}@}}",
            "[[": f"[[{var_name}]]",
            "(@": f"(@{var_name}@)"
        }
        
        variable = delim_map.get(delimiter, f"{{{{{var_name}}}}}")
        self.replace_selection(variable)
        
        self.custom_var_entry.delete(0, "end")
    
    def replace_selection(self, variable):
        """Replace selected text with variable"""
        try:
            start_index = self.text_widget.index("sel.first")
            end_index = self.text_widget.index("sel.last")
            selected_text = self.text_widget.get(start_index, end_index)
            
            # Record replacement for undo
            self.replacements.append({
                "start": start_index,
                "end": end_index,
                "original": selected_text,
                "variable": variable
            })
            
            # Replace text
            self.text_widget.delete(start_index, end_index)
            self.text_widget.insert(start_index, variable)
            
            # Highlight the variable
            var_end = self.text_widget.index(f"{start_index}+{len(variable)}c")
            self.text_widget.tag_add("variable", start_index, var_end)
            
            # Log the replacement
            self.log_message(f"'{selected_text[:30]}...' → {variable}")
            
        except tk.TclError:
            messagebox.showwarning("No Selection", "Please select text first.")
    
    def undo_last(self):
        """Undo the last replacement"""
        if not self.replacements:
            messagebox.showinfo("Nothing to Undo", "No replacements to undo.")
            return
        
        last = self.replacements.pop()
        
        # Find current position of the variable
        # This is approximate since text may have shifted
        try:
            # Remove the variable
            start = last["start"]
            var_len = len(last["variable"])
            end = self.text_widget.index(f"{start}+{var_len}c")
            
            self.text_widget.delete(start, end)
            self.text_widget.insert(start, last["original"])
            
            self.log_message(f"↩️ UNDONE: {last['variable']} → '{last['original'][:30]}'")
            
        except Exception as e:
            messagebox.showerror("Undo Error", f"Failed to undo:\n{e}")
    
    def log_message(self, message):
        """Add message to replacements log"""
        self.log_text.config(state="normal")
        self.log_text.insert("end", f"{message}\n")
        self.log_text.see("end")
        self.log_text.config(state="disabled")
    
    def save_template(self):
        """Save the modified document as a template"""
        if not self.doc:
            messagebox.showwarning("No Document", "Please load a document first.")
            return
        
        # Get modified text
        modified_text = self.text_widget.get("1.0", "end-1c")
        
        # Ask for save location
        save_path = filedialog.asksaveasfilename(
            title="Save Template As",
            defaultextension=".docx",
            filetypes=[("Word Template", "*.docx")],
            initialfile=f"{self.doc_path.stem}_template.docx" if self.doc_path else "template.docx"
        )
        
        if not save_path:
            return
        
        try:
            # Replace paragraphs in original document
            lines = modified_text.split("\n")
            
            # Clear existing paragraphs
            for paragraph in self.doc.paragraphs:
                paragraph.clear()
            
            # Add new content
            for i, line in enumerate(lines):
                if i < len(self.doc.paragraphs):
                    self.doc.paragraphs[i].add_run(line)
                else:
                    self.doc.add_paragraph(line)
            
            # Save
            self.doc.save(save_path)
            
            messagebox.showinfo(
                "Template Saved",
                f"Template saved successfully!\n\n"
                f"Location: {save_path}\n\n"
                f"Replacements made: {len(self.replacements)}"
            )
            
            self.log_message(f"✓ Saved template: {Path(save_path).name}")
            
        except Exception as e:
            messagebox.showerror("Save Error", f"Failed to save template:\n{e}")
    
    def run(self):
        """Start the template builder"""
        self.window.mainloop()


def open_template_builder(parent=None):
    """Open the template builder tool"""
    builder = TemplateBuilder(parent)
    builder.run()